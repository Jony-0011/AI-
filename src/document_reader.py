#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档读取模块
负责读取和解析docx/pdf文档，提取文本和图片信息
"""

import os
from pathlib import Path
from typing import List, Dict, Any
from docx import Document


class DocumentReader:
    """文档读取器类，处理文档读取与解析"""

    def __init__(self, file_path: str = None):
        self.file_path = file_path
        self.text_content = ""
        self.images = []
        self.doc_type = None
        if file_path:
            self._validate_file()

    def _validate_file(self):
        """验证文件是否存在和格式是否支持"""
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"文件不存在: {self.file_path}")

        ext = Path(self.file_path).suffix.lower()
        if ext not in ['.docx']:
            raise ValueError(f"不支持的文件格式: {ext}，仅支持.docx")
        self.doc_type = ext

    def read_docx(self) -> Dict[str, Any]:
        """读取docx文档，提取文本和图片信息"""
        doc = Document(self.file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        self.text_content = '\n'.join(text_parts)

        self._extract_docx_images(doc)

        return {
            'text': self.text_content,
            'images': self.images,
            'file_path': self.file_path,
            'file_type': self.doc_type
        }

    def _extract_docx_images(self, doc):
        """从docx中提取图片信息"""
        images_found = set()
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref.lower():
                image_info = {
                    'format': rel.target_ref.split('.')[-1].lower(),
                    'rid': rel.rId
                }

                try:
                    image_part = rel.target_part
                    image_info['size_bytes'] = len(image_part.blob)
                except:
                    image_info['size_bytes'] = 0

                if rel.rId not in images_found:
                    self.images.append(image_info)
                    images_found.add(rel.rId)
        
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run._element.xpath('.//pic:pic'):
                    if 'inline_image' not in images_found:
                        self.images.append({
                            'format': 'png',
                            'type': 'inline',
                            'rid': f'inline_{len(self.images)}'
                        })
                        images_found.add('inline_image')

    def read(self) -> Dict[str, Any]:
        """统一的读取入口方法"""
        if self.doc_type == '.docx':
            return self.read_docx()
        else:
            raise ValueError(f"未知的文档类型: {self.doc_type}")

    def get_structure(self) -> List[str]:
        """识别文档结构（章节标题等）"""
        structure_keywords = ['一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、',
                              '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.',
                              '实验目的', '实验原理', '实验步骤', '实验结果', '实验反思', '实验总结'
                             ]

        sections = []
        lines = self.text_content.split('\n')
        for line in lines:
            line = line.strip()
            for keyword in structure_keywords:
                if keyword in line and len(line) < 50:
                    if line not in sections:
                        sections.append(line)
        return sections

    def read_docx_from_bytes(self, file_bytes: bytes) -> Dict[str, Any]:
        """从字节流读取docx文档，提取文本和图片信息"""
        from io import BytesIO
        
        doc = Document(BytesIO(file_bytes))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        self.text_content = '\n'.join(text_parts)
        self.images = []
        self._extract_docx_images(doc)

        return {
            'text': self.text_content,
            'images': self.images,
            'file_type': '.docx'
        }
